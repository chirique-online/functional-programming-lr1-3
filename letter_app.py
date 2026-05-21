import tkinter as tk
from tkinter import messagebox, filedialog, ttk
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
import shutil
from datetime import datetime

class AppendixDialog:
    def __init__(self, parent, title="", text=""):
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Приложение")
        self.dialog.geometry("550x450")
        self.dialog.transient(parent)
        self.dialog.grab_set()

        self.result = None

        tk.Label(self.dialog, text="Заголовок приложения (будет по центру):").pack(pady=5)
        self.title_entry = tk.Entry(self.dialog, width=60)
        self.title_entry.insert(0, title)
        self.title_entry.pack()

        tk.Label(self.dialog, text="Текст приложения:").pack(pady=5)
        self.text_text = tk.Text(self.dialog, height=15, width=60)
        self.text_text.insert("1.0", text)
        self.text_text.pack()

        btn_frame = tk.Frame(self.dialog)
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="OK", command=self.ok).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="Отмена", command=self.cancel).pack(side=tk.LEFT, padx=5)

    def ok(self):
        title = self.title_entry.get().strip()
        text = self.text_text.get("1.0", tk.END).strip()
        if not title:
            messagebox.showwarning("Ошибка", "Заголовок не может быть пустым")
            return
        self.result = {"title": title, "text": text}
        self.dialog.destroy()

    def cancel(self):
        self.dialog.destroy()


class LetterWithAppendices:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор писем с приложениями - Лабораторная №3")
        self.root.geometry("1400x800")
        self.root.resizable(True, True)

        script_dir = os.path.dirname(os.path.abspath(__file__))
        default_template = os.path.join(script_dir, "шаблон_с_приложениями.docx")
        
        if os.path.exists(default_template):
            self.template_path = default_template
        else:
            self.template_path = ""
        
        self.appendices = []

        self.create_widgets()
        
        if self.template_path:
            self.template_label.config(text=f"Шаблон: {os.path.basename(self.template_path)}", fg="green")
            self.generate_btn.config(state=tk.NORMAL)
        
        self.update_preview()

    def create_widgets(self):
        main_paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        left_frame = ttk.Frame(main_paned)
        main_paned.add(left_frame, weight=1)

        left_canvas = tk.Canvas(left_frame)
        left_scrollbar = tk.Scrollbar(left_frame, orient="vertical", command=left_canvas.yview)
        left_scrollable = tk.Frame(left_canvas)

        left_scrollable.bind("<Configure>", lambda e: left_canvas.configure(scrollregion=left_canvas.bbox("all")))
        left_canvas.create_window((0, 0), window=left_scrollable, anchor="nw")
        left_canvas.configure(yscrollcommand=left_scrollbar.set)

        left_canvas.pack(side="left", fill="both", expand=True)
        left_scrollbar.pack(side="right", fill="y")

        self.entries = {}
        row = 0

        sender_frame = ttk.LabelFrame(left_scrollable, text="📤 ОТПРАВИТЕЛЬ", padding=10)
        sender_frame.grid(row=row, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        row += 1

        sender_fields = [
            ("Название организации:", "org_name", "АО «Пролетарский прогресс»"),
            ("Адрес:", "org_address", "г. Москва, ул. Ленина, д. 5"),
            ("Телефон:", "org_phone", "(495) 123-45-67"),
            ("E-mail:", "org_email", "info@example.ru"),
        ]

        for i, (label_text, key, default) in enumerate(sender_fields):
            tk.Label(sender_frame, text=label_text).grid(row=i, column=0, sticky="w", pady=2, padx=5)
            entry = tk.Entry(sender_frame, width=50)
            entry.insert(0, default)
            entry.grid(row=i, column=1, pady=2, padx=5)
            entry.bind("<KeyRelease>", lambda e: self.update_preview())
            self.entries[key] = entry

        req_frame = ttk.LabelFrame(left_scrollable, text="📄 РЕКВИЗИТЫ ПИСЬМА", padding=10)
        req_frame.grid(row=row, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        row += 1

        tk.Label(req_frame, text="Дата (дд.мм.гггг):").grid(row=0, column=0, sticky="w", pady=2, padx=5)
        self.date_entry = tk.Entry(req_frame, width=30)
        self.date_entry.insert(0, datetime.now().strftime("%d.%m.%Y"))
        self.date_entry.grid(row=0, column=1, pady=2, padx=5)
        self.date_entry.bind("<KeyRelease>", lambda e: self.update_preview())

        tk.Label(req_frame, text="Исходящий номер:").grid(row=1, column=0, sticky="w", pady=2, padx=5)
        self.out_number_entry = tk.Entry(req_frame, width=30)
        self.out_number_entry.insert(0, "12/05")
        self.out_number_entry.grid(row=1, column=1, pady=2, padx=5)
        self.out_number_entry.bind("<KeyRelease>", lambda e: self.update_preview())

        tk.Label(req_frame, text="Входящий номер (если есть):").grid(row=2, column=0, sticky="w", pady=2, padx=5)
        self.in_number_entry = tk.Entry(req_frame, width=30)
        self.in_number_entry.insert(0, "45")
        self.in_number_entry.grid(row=2, column=1, pady=2, padx=5)
        self.in_number_entry.bind("<KeyRelease>", lambda e: self.update_preview())

        recipient_frame = ttk.LabelFrame(left_scrollable, text="📥 ПОЛУЧАТЕЛЬ", padding=10)
        recipient_frame.grid(row=row, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        row += 1

        recipient_fields = [
            ("Должность адресата (в дательном падеже):", "recipient_position", "Генеральному директору"),
            ("Организация адресата:", "recipient_company", "ООО «Недвижимость»"),
            ("ФИО адресата (в дательном падеже):", "recipient_name", "Иванову И.И."),
        ]

        for i, (label_text, key, default) in enumerate(recipient_fields):
            tk.Label(recipient_frame, text=label_text).grid(row=i, column=0, sticky="w", pady=2, padx=5)
            entry = tk.Entry(recipient_frame, width=50)
            entry.insert(0, default)
            entry.grid(row=i, column=1, pady=2, padx=5)
            entry.bind("<KeyRelease>", lambda e: self.update_preview())
            self.entries[key] = entry

        body_frame = ttk.LabelFrame(left_scrollable, text="ТЕКСТ ПИСЬМА", padding=10)
        body_frame.grid(row=row, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        row += 1

        tk.Label(body_frame, text="Содержание (можно использовать {RECIPIENT_NAME}):").pack(anchor="w", pady=2)
        default_body = """Уважаемому {RECIPIENT_NAME}!

Настоящим письмом направляем вам информацию и прилагаем необходимые документы.

Просим рассмотреть и сообщить о решении."""
        self.body_text = tk.Text(body_frame, height=8, width=55)
        self.body_text.insert("1.0", default_body)
        self.body_text.pack(pady=5)
        self.body_text.bind("<KeyRelease>", lambda e: self.update_preview())

        appendix_frame = ttk.LabelFrame(left_scrollable, text="📎 ПРИЛОЖЕНИЯ", padding=10)
        appendix_frame.grid(row=row, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        row += 1

        self.appendices_listbox = tk.Listbox(appendix_frame, height=5, width=70)
        self.appendices_listbox.pack(pady=5)
        self.appendices_listbox.bind("<<ListboxSelect>>", lambda e: self.update_preview())

        btn_app_frame = tk.Frame(appendix_frame)
        btn_app_frame.pack(pady=5)
        tk.Button(btn_app_frame, text="Добавить", command=self.add_appendix).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_app_frame, text="Редактировать", command=self.edit_appendix).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_app_frame, text="Удалить", command=self.delete_appendix).pack(side=tk.LEFT, padx=5)

        sign_frame = ttk.LabelFrame(left_scrollable, text="✍️ ПОДПИСЬ И ИСПОЛНИТЕЛЬ", padding=10)
        sign_frame.grid(row=row, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        row += 1

        sign_fields = [
            ("Должность подписанта:", "signer_position", "Генеральный директор"),
            ("ФИО подписанта:", "signer_name", "Петров П.П."),
            ("Исполнитель (ФИО):", "executor_name", "Сидорова А.А."),
            ("Телефон исполнителя:", "executor_phone", "(495) 999-88-77"),
        ]

        for i, (label_text, key, default) in enumerate(sign_fields):
            tk.Label(sign_frame, text=label_text).grid(row=i, column=0, sticky="w", pady=2, padx=5)
            entry = tk.Entry(sign_frame, width=50)
            entry.insert(0, default)
            entry.grid(row=i, column=1, pady=2, padx=5)
            entry.bind("<KeyRelease>", lambda e: self.update_preview())
            self.entries[key] = entry

        control_frame = ttk.LabelFrame(left_scrollable, text="УПРАВЛЕНИЕ", padding=10)
        control_frame.grid(row=row, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        row += 1

        self.template_btn = tk.Button(control_frame, text="Выбрать другой шаблон DOCX", command=self.select_template, bg="lightgray")
        self.template_btn.pack(pady=5)

        self.template_label = tk.Label(control_frame, text="Шаблон не выбран", fg="red")
        self.template_label.pack()

        self.generate_btn = tk.Button(control_frame, text="Создать DOCX-письмо", command=self.generate_letter, bg="lightblue", state=tk.DISABLED)
        self.generate_btn.pack(pady=5)

        self.status_label = tk.Label(control_frame, text="Готов", fg="green")
        self.status_label.pack()

        right_frame = ttk.Frame(main_paned)
        main_paned.add(right_frame, weight=1)

        preview_label = tk.Label(right_frame, text="🔍 ПРЕДПРОСМОТР ПИСЬМА", font=("Arial", 10, "bold"))
        preview_label.pack(pady=5)

        self.preview_text = tk.Text(right_frame, wrap=tk.WORD, font=("Courier", 10))
        self.preview_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        preview_scrollbar = tk.Scrollbar(self.preview_text)
        preview_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.preview_text.config(yscrollcommand=preview_scrollbar.set)
        preview_scrollbar.config(command=self.preview_text.yview)

    def get_template_structure(self):
        if not self.template_path or not os.path.exists(self.template_path):
            return None
        
        try:
            doc = Document(self.template_path)
            template_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    template_text.append(paragraph.text)
            return "\n".join(template_text)
        except:
            return None

    def update_preview(self):
        data = {
            "org_name": self.entries["org_name"].get().strip(),
            "org_address": self.entries["org_address"].get().strip(),
            "org_phone": self.entries["org_phone"].get().strip(),
            "org_email": self.entries["org_email"].get().strip(),
            "date": self.date_entry.get().strip(),
            "out_number": self.out_number_entry.get().strip(),
            "in_number": self.in_number_entry.get().strip(),
            "recipient_position": self.entries["recipient_position"].get().strip(),
            "recipient_company": self.entries["recipient_company"].get().strip(),
            "recipient_name": self.entries["recipient_name"].get().strip(),
            "signer_position": self.entries["signer_position"].get().strip(),
            "signer_name": self.entries["signer_name"].get().strip(),
            "executor_name": self.entries["executor_name"].get().strip(),
            "executor_phone": self.entries["executor_phone"].get().strip(),
        }

        body_text = self.body_text.get("1.0", tk.END).strip()
        body_text = body_text.replace("{RECIPIENT_NAME}", data["recipient_name"])

        preview_lines = []
        
        template_structure = self.get_template_structure()
        
        if template_structure:
            preview = template_structure
            preview = preview.replace("{ORG_NAME}", data["org_name"])
            preview = preview.replace("{ORG_ADDRESS}", data["org_address"])
            preview = preview.replace("{ORG_PHONE}", data["org_phone"])
            preview = preview.replace("{ORG_EMAIL}", data["org_email"])
            preview = preview.replace("{DATE}", data["date"])
            preview = preview.replace("{OUT_NUMBER}", data["out_number"])
            preview = preview.replace("{IN_NUMBER}", data["in_number"])
            preview = preview.replace("{RECIPIENT_POSITION}", data["recipient_position"])
            preview = preview.replace("{RECIPIENT_COMPANY}", data["recipient_company"])
            preview = preview.replace("{RECIPIENT_NAME}", data["recipient_name"])
            preview = preview.replace("{SIGNER_POSITION}", data["signer_position"])
            preview = preview.replace("{SIGNER_NAME}", data["signer_name"])
            preview = preview.replace("{EXECUTOR_NAME}", data["executor_name"])
            preview = preview.replace("{EXECUTOR_PHONE}", data["executor_phone"])
            preview = preview.replace("{BODY}", body_text)
            
            if self.appendices:
                app_list = "Приложение:\n" + "\n".join([f"{idx}. {app['title']}" for idx, app in enumerate(self.appendices, 1)])
                preview = preview.replace("{APPENDICES_LIST}", app_list)
            else:
                preview = preview.replace("{APPENDICES_LIST}", "")
            
            preview = preview.replace("{APPENDICES_CONTENT}", "\n[Страницы приложений будут сгенерированы в DOCX]")
            
            preview_lines = preview.split("\n")
        else:
            preview_lines.append(data["org_name"])
            preview_lines.append(data["org_address"])
            preview_lines.append(f"тел. {data['org_phone']}   E-mail: {data['org_email']}")
            preview_lines.append("")
            preview_lines.append(data["recipient_position"])
            preview_lines.append(data["recipient_company"])
            preview_lines.append(data["recipient_name"])
            preview_lines.append("")
            preview_lines.append(f"{data['date']}   № {data['out_number']}")
            preview_lines.append(f"На № {data['in_number']} от __________")
            preview_lines.append("")
            preview_lines.append(body_text)
            preview_lines.append("")

            if self.appendices:
                preview_lines.append("Приложение:")
                for idx, app in enumerate(self.appendices, 1):
                    preview_lines.append(f"{idx}. {app['title']}")
                preview_lines.append("")

            preview_lines.append("С уважением,")
            preview_lines.append("")
            preview_lines.append(f"{data['signer_position']}                {data['signer_name']}")
            preview_lines.append("")
            preview_lines.append(f"Исполнитель: {data['executor_name']}, тел. {data['executor_phone']}")

        self.preview_text.delete("1.0", tk.END)
        self.preview_text.insert("1.0", "\n".join(preview_lines))

    def update_appendix_listbox(self):
        self.appendices_listbox.delete(0, tk.END)
        for idx, app in enumerate(self.appendices, 1):
            self.appendices_listbox.insert(tk.END, f"{idx}. {app['title']}")
        self.update_preview()

    def add_appendix(self):
        dialog = AppendixDialog(self.root)
        self.root.wait_window(dialog.dialog)
        if dialog.result:
            self.appendices.append(dialog.result)
            self.update_appendix_listbox()

    def edit_appendix(self):
        selection = self.appendices_listbox.curselection()
        if not selection:
            messagebox.showinfo("Редактирование", "Выберите приложение для редактирования")
            return
        idx = selection[0]
        app = self.appendices[idx]
        dialog = AppendixDialog(self.root, title=app["title"], text=app["text"])
        self.root.wait_window(dialog.dialog)
        if dialog.result:
            self.appendices[idx] = dialog.result
            self.update_appendix_listbox()

    def delete_appendix(self):
        selection = self.appendices_listbox.curselection()
        if not selection:
            return
        idx = selection[0]
        del self.appendices[idx]
        self.update_appendix_listbox()

    def select_template(self):
        file_path = filedialog.askopenfilename(filetypes=[("DOCX файлы", "*.docx")])
        if file_path:
            self.template_path = file_path
            self.template_label.config(text=f"Шаблон: {os.path.basename(file_path)}", fg="green")
            self.generate_btn.config(state=tk.NORMAL)
            self.update_preview() 

    def replace_placeholders(self, doc, replacements):
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, replacements)

    def replace_in_paragraph(self, paragraph, replacements):
        full_text = paragraph.text
        new_text = full_text
        for placeholder, value in replacements.items():
            if value is None:
                value = ""
            new_text = new_text.replace(placeholder, value)
        if new_text != full_text:
            paragraph.clear()
            paragraph.add_run(new_text)

    def generate_appendix_list_text(self):
        if not self.appendices:
            return ""
        lines = ["Приложение:"]
        for idx, app in enumerate(self.appendices, 1):
            lines.append(f"{idx}. {app['title']}")
        return "\n".join(lines)

    def insert_appendices_into_doc(self, doc):
        for paragraph in doc.paragraphs:
            if "{APPENDICES_CONTENT}" in paragraph.text:
                p = paragraph._element
                p.getparent().remove(p)
                break

        if not self.appendices:
            return

        doc.add_page_break()

        for idx, app in enumerate(self.appendices, 1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if len(self.appendices) == 1:
                p.add_run("Приложение")
            else:
                p.add_run(f"Приложение {idx}")

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(app["title"])
            run.bold = True
            run.font.size = Pt(14)

            lines = app["text"].splitlines()
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()

            if idx < len(self.appendices):
                doc.add_page_break()

    def validate_fields(self):
        errors = []

        if not self.entries["org_name"].get().strip():
            errors.append("Название организации-отправителя")

        if not self.entries["recipient_name"].get().strip():
            errors.append("ФИО получателя")

        body_text = self.body_text.get("1.0", tk.END).strip()
        if not body_text:
            errors.append("Текст письма")

        if errors:
            messagebox.showwarning(
                "Незаполненные поля",
                f"Пожалуйста, заполните следующие обязательные поля:\n" + "\n".join(f"• {e}" for e in errors)
            )
            return False
        return True

    def generate_letter(self):
        if not self.validate_fields():
            return

        if not self.template_path:
            messagebox.showerror("Ошибка", "Сначала выберите шаблон!")
            return

        replacements = {
            "{ORG_NAME}": self.entries["org_name"].get().strip(),
            "{ORG_ADDRESS}": self.entries["org_address"].get().strip(),
            "{ORG_PHONE}": self.entries["org_phone"].get().strip(),
            "{ORG_EMAIL}": self.entries["org_email"].get().strip(),
            "{DATE}": self.date_entry.get().strip(),
            "{OUT_NUMBER}": self.out_number_entry.get().strip(),
            "{IN_NUMBER}": self.in_number_entry.get().strip(),
            "{RECIPIENT_POSITION}": self.entries["recipient_position"].get().strip(),
            "{RECIPIENT_COMPANY}": self.entries["recipient_company"].get().strip(),
            "{RECIPIENT_NAME}": self.entries["recipient_name"].get().strip(),
            "{SIGNER_POSITION}": self.entries["signer_position"].get().strip(),
            "{SIGNER_NAME}": self.entries["signer_name"].get().strip(),
            "{EXECUTOR_NAME}": self.entries["executor_name"].get().strip(),
            "{EXECUTOR_PHONE}": self.entries["executor_phone"].get().strip(),
        }

        body_text = self.body_text.get("1.0", tk.END).strip()
        body_text = body_text.replace("{RECIPIENT_NAME}", replacements["{RECIPIENT_NAME}"])
        replacements["{BODY}"] = body_text

        replacements["{APPENDICES_LIST}"] = self.generate_appendix_list_text()

        output_dir = os.path.join(os.path.expanduser("~"), "Desktop")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(output_dir, f"Letter_{timestamp}.docx")

        try:
            shutil.copy2(self.template_path, output_path)
            doc = Document(output_path)

            self.replace_placeholders(doc, replacements)
            self.insert_appendices_into_doc(doc)

            doc.save(output_path)

            self.status_label.config(text=f"Готово: {output_path}", fg="blue")
            messagebox.showinfo("Успех", f"Письмо создано!\n{output_path}")
        except Exception as e:
            self.status_label.config(text="Ошибка!", fg="red")
            messagebox.showerror("Ошибка", str(e))


if __name__ == "__main__":
    root = tk.Tk()
    app = LetterWithAppendices(root)
    root.mainloop()